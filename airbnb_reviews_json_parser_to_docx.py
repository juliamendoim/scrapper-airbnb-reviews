import json
import os
import codecs
from docx import Document
 
datos_path = "path/reviews/"


 #Loopeo por cada archivo
 #data es un documento con 7 comentarios (adentro de reviews)
for file in os.listdir(datos_path):
	data_file = open(datos_path+file)
	data = json.load(data_file)

	i = 0

	
#Extrae cada comentario, lenguaje en el que está el comentario y id del usuario que comentó. Los escribe en un doc.	
	documento = Document()
	
	for elemento in data["reviews"]:

		x = (data["reviews"][i]["comments"])
		y = (data["reviews"][i]["language"])
		z = (data["reviews"][i]["id"])
		
		documento.add_paragraph(str(x))
		documento.add_paragraph(detect_lan(str(y)))
		documento.add_paragraph(str(z))
		
		i += 1
	documento.save(str('reviewsdocx/'+str(file[0:-5])+'.docx'))

#O, se pueden volcar en un .txt (acá, filtra por idioma) - descomentar
#if y == 'es':
#	with codecs.open('/path/airbnb/txt/'+str(file[0:-5])+str(i)+'.txt','w', 'utf-8') as f:
#		f.write(str(x))



	